from __future__ import annotations

import io

import pdfplumber
from docx import Document

from app.utils.logging import get_logger

logger = get_logger(__name__)


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from a PDF using pdfplumber."""
    text_parts: list[str] = []

    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        logger.info("PDF extraction started", extra={"page_count": len(pdf.pages)})
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
            if page_text:
                logger.debug(
                    "Page extracted",
                    extra={"page": i + 1, "chars": len(page_text)},
                )

    result = "\n".join(text_parts)
    logger.info(
        "PDF extraction complete",
        extra={"total_chars": len(result), "pages_with_text": sum(1 for t in text_parts if t)},
    )
    return result


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from a DOCX file including paragraphs and tables."""
    doc = Document(io.BytesIO(file_content))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    result = "\n".join(parts)
    logger.info(
        "DOCX extraction complete",
        extra={
            "total_chars": len(result),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        },
    )
    return result
